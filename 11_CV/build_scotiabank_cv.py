from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins
section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

heading_color = RGBColor(0x1F, 0x4E, 0x79)

def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = heading_color
    run.font.name = 'Calibri'
    p.space_after = Pt(2)

def add_line(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    p.space_after = Pt(2)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = heading_color
    run.font.name = 'Calibri'

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)

def add_plain(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)

# Header
add_name('Felipe Maldonado')
add_line('Senior Manager Data & Insights | BI Strategy | Executive Reporting | Financial Services Analytics', bold=True, size=11)
add_line('Montreal, QC  |  +1 438 340 3252  |  inq.felipe.maldonado@gmail.com  |  linkedin.com/in/felipe-maldonado')

# Professional Profile
add_section_heading('Professional Profile')
add_plain(
    'Senior Data & Insights professional with 12+ years of experience transforming business data into strategic insights '
    'that support executive decision-making, business growth, and operational excellence. Proven ability to partner with '
    'senior leaders and cross-functional teams across financial services, banking, sales, marketing, and operations to define '
    'KPIs, develop executive reporting, and enable data-driven strategies. Strong background in financial services analytics, '
    'including banking campaigns, customer engagement, commercial performance, CRM analytics, and enterprise reporting. Trilingual '
    'professional (Spanish, English, French) with experience collaborating with business executives, technical teams, and '
    'international stakeholders.'
)

# Core Competencies
add_section_heading('Core Competencies')
competencies = [
    'Executive reporting & data storytelling',
    'KPI / OKR development and performance management',
    'Financial services, banking, and campaign analytics',
    'Data aggregation, reconciliation, and quality assurance',
    'Dashboard development (Power BI / Tableau)',
    'Advanced SQL, data modeling, and ETL/ELT',
    'Stakeholder management & cross-functional leadership',
    'Change management & continuous improvement',
    'AI tools awareness and adoption',
    'Project delivery & requirements management',
]
for c in competencies:
    add_bullet(c)

# Helper for experience roles
def add_role(company, title, location, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{title}  |  {company}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(f'{dates}  |  {location}')
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    for b in bullets:
        add_bullet(b)

# Professional Experience
add_section_heading('Professional Experience')

add_role(
    'Alithya',
    'Senior Business Intelligence Analyst',
    'Montreal, QC',
    'Aug 2022 – Present',
    [
        'Partner with business leaders and stakeholders to transform business questions into data-driven insights, reporting strategies, and analytical solutions supporting financial and operational decision-making.',
        'Define KPIs, reporting strategies, and business performance indicators for senior stakeholders.',
        'Develop executive dashboards and reporting solutions using Power BI, SQL, Snowflake, and Microsoft Fabric.',
        'Consolidate and analyze data from multiple business domains to identify trends, opportunities, and operational improvements.',
        'Investigate data discrepancies, reconcile information sources, and ensure reporting accuracy and reliability.',
        'Manage analytics initiatives from requirements gathering through implementation, coordinating business and technical teams.',
        'Support change management initiatives by improving adoption of self-service analytics and data-driven processes.',
    ]
)

add_role(
    'CGI',
    'Data Analyst',
    'Montreal, QC',
    'Jul 2021 – Aug 2022',
    [
        'Partnered with business stakeholders to translate strategic requirements into logical and physical data models, improving data quality and supporting long-term enterprise solutions.',
        'Supported the operationalization of data governance policies, standards, and procedures.',
        'Led cross-functional project coordination independently, bridging technical and business teams to deliver analytics solutions on time.',
        'Designed and deployed self-service reporting capabilities and user enablement documentation.',
        'Designed, implemented, and optimized ETL pipelines using Talend and Pentaho.',
    ]
)

add_role(
    'AECSA – Lawyers specialized in bank collection',
    'Strategic Data Director (BI)',
    'Bogotá, Colombia',
    '2017 – 2018',
    [
        'Designed and developed database specifications for dashboards and data visualization for key projects.',
        'Developed and optimized ETL/SQL programs to improve runtime efficiency and reporting performance.',
        'Monitored business operations and provided data-driven recommendations to support decision-making.',
        'Led automation and process reengineering initiatives that improved data storage, integration, and exploitation.',
        'Identified analytics opportunities that improved sales department performance and strategic planning.',
    ]
)

add_role(
    'Banco Falabella Colombia',
    'Business Intelligence Campaign Coordinator',
    'Bogotá, Colombia',
    'Dec 2014 – Nov 2016',
    [
        'Gathered and managed business requirements from marketing, sales, and commercial stakeholders to design analytics solutions supporting campaign planning and execution.',
        'Designed database structures and data models to support campaign analytics, reporting systems, and performance tracking.',
        'Created KPI frameworks and executive reporting capabilities to monitor campaign performance, customer behavior, and business impact.',
        'Improved data quality processes and automated business reporting to increase reliability and reduce manual effort.',
    ]
)

add_role(
    'SoftTek Solutions - INFOMEDIA',
    'IT Project Manager – Business Intelligence',
    'Mexico City, Mexico',
    'Aug 2013 – Sep 2014',
    [
        'Led the implementation and adoption of BI tools, systems, processes, and indicators across cross-functional teams.',
        'Designed and developed ETL processes to integrate and consolidate data from heterogeneous systems.',
        'Verified data quality and accuracy across multiple information systems before executive reporting.',
        'Acted as a change management leader, presenting observations and recommendations to management.',
    ]
)

add_role(
    'BBVA Colombia',
    'BI Specialist Professional',
    'Bogotá, Colombia',
    'Jan 2008 – Jul 2013',
    [
        'Developed banking applications and information exploitation procedures to support customer analytics and commercial intelligence.',
        'Participated in corporate BI projects across Latin America, representing BBVA Colombia during the rollout of the Campaign Manager and Commercial Agenda project.',
        'Designed reporting solutions to support customer engagement, campaign management, and business performance monitoring.',
        'Provided analytical insights supporting commercial strategy and digital channel improvements.',
    ]
)

# Additional Experience
add_section_heading('Additional Experience')
add_plain('Freelance — Data Analyst (BI)  |  Montreal, QC  |  2019 – 2020')
add_plain('oXya, A Hitachi Group Company — Operations Technician  |  Montreal, QC  |  2021')
add_plain('InteRecycle.com — Technical Support Specialist  |  Montreal, QC  |  2020')

# Technical Skills
add_section_heading('Technical Skills')
skills = [
    'BI & Analytics: Power BI, Power BI Service, DAX, Power Query, Tableau, Salesforce Analytics, Executive Dashboards, Self-Service BI',
    'Data & Engineering: SQL, Snowflake, Microsoft Fabric, Microsoft SQL Server, Data Warehousing, ETL/ELT, Data Modeling, Python',
    'Data Quality & Governance: Data Validation, Data Reconciliation, Data Governance, Data Lineage, Reporting Accuracy Controls',
    'CRM & Financial Analytics: Salesforce Marketing Cloud, Campaign Analytics, Customer Segmentation, Banking Analytics',
    'Collaboration & Delivery: Jira, Confluence, Git, Agile / Scrum / Kanban',
    'Leadership: Executive Stakeholder Management, Data Storytelling, Executive Presentation, Cross-functional Leadership, Change Management',
]
for s in skills:
    add_bullet(s)

# Education
add_section_heading('Education')
add_plain('Master’s in Big Data and Business Intelligence — ENEB - Barcelona European Business School — 2021 – In Progress')
add_plain('Diploma of Vocational Studies – Computing Support — Teccart Institute, Montreal, QC — 2018 – 2020')
add_plain('Master’s (Specialization) in Digital Marketing — IEP - European University Institute, Madrid, Spain — 2014 – 2015')
add_plain('Bachelor of Systems Engineering — University of Cundinamarca — 2002 – 2007')

# Languages
add_section_heading('Languages')
add_plain('Spanish — Native  |  English — Professional working proficiency  |  French — Professional working proficiency')

# Save
output_path = r'e:\bi-data-engineering-knowledge-base-main\11_CV\CV_Felipe_Maldonado_Senior_Manager_Data_and_Insights_Scotiabank.docx'
doc.save(output_path)
print(f'DOCX created: {output_path}')
