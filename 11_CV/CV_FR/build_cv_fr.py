from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Marges
section = doc.sections[0]
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

# Police par défaut
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

heading_color = RGBColor(0x1F, 0x4E, 0x79)

def add_name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = heading_color
    run.font.name = 'Calibri'
    p.space_after = Pt(2)

def add_line(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    p.space_after = Pt(2)

def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = heading_color
    run.font.name = 'Calibri'

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    for run in p.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(1)

def add_plain(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)

# En-tête
add_name('Felipe Maldonado')
add_line('Leader Sénior en Business Intelligence et Données | Stratégie BI | Reporting Exécutif | Analytique des Services Financiers', bold=True, size=11)
add_line('Montréal, QC  |  +1 438 340 3252  |  inq.felipe.maldonado@gmail.com  |  linkedin.com/in/felipe-maldonado')

# Profil Professionnel
add_section_heading('Profil Professionnel')
add_plain(
    'Professionnel sénior en business intelligence et en données avec 18 ans d\u2019expérience, des systèmes transactionnels '
    'mainframe jusqu\u2019aux plateformes de données infonuagiques modernes, dont 11 ans dans les services financiers (banque de '
    'détail, crédit, gestion de créances et LBA/AML) au Canada, en Colombie, au Mexique et en Espagne. Je dirige actuellement le '
    'volet BI d\u2019une migration stratégique d\u2019un système opérationnel sur site vers l\u2019infonuagique, avec une architecture cible '
    'dans Snowflake certifiée à 99 % d\u2019exactitude au niveau de l\u2019enregistrement, et j\u2019ai conçu de bout en bout une plateforme '
    'd\u2019entreprise de commissions de vente. J\u2019ai géré des équipes allant jusqu\u2019à six personnes et standardisé les définitions de '
    'KPI sur 30 portefeuilles clients et quatre secteurs d\u2019affaires. Trilingue (espagnol, français, anglais), à l\u2019aise tant avec '
    'les intervenants exécutifs qu\u2019avec les équipes techniques.'
)

# Compétences Clés
add_section_heading('Compétences Clés')
competencies = [
    'Stratégie de business intelligence et reporting exécutif',
    'Cadres de KPI / OKR et gestion de la performance d\u2019affaires',
    'Analytique des services financiers, bancaires et de campagnes',
    'Réconciliation, validation et gouvernance des données',
    'Développement de tableaux de bord (Power BI)',
    'SQL avancé, modélisation de données et ETL/ELT',
    'Gestion d\u2019intervenants exécutifs et leadership interfonctionnel',
    'Gestion du changement et amélioration continue',
    'Ingénierie assistée par IA (Devin, Windsurf / Cascade)',
    'Livraison de projets et gestion des exigences d\u2019affaires',
]
for c in competencies:
    add_bullet(c)

# Fonction pour les rôles d'expérience
def add_role(company, title, location, dates, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{title}  |  {company}')
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    run2 = p2.add_run(f'{dates}  |  {location}')
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    for b in bullets:
        add_bullet(b)

# Expérience Professionnelle
add_section_heading('Expérience Professionnelle')

add_role(
    'Alithya',
    'Senior Business Intelligence Analyst',
    'Montréal, QC',
    'Août 2022 - Présent',
    [
        'J\u2019ai défini l\u2019architecture cible de données dans Snowflake et dirigé le volet BI de l\u2019analyse des systèmes sources jusqu\u2019à l\u2019approbation en production.',
        'J\u2019ai bâti la validation massive d\u2019intégrité qui a certifié la migration à 99 % d\u2019exactitude au niveau de l\u2019enregistrement sur 5 cycles de validation, le seuil fixé par l\u2019entreprise pour approuver la mise en production.',
        'J\u2019ai conçu et bâti de bout en bout une plateforme d\u2019entreprise de commissions de vente en tant que seul responsable (analyse d\u2019affaires, modélisation SCD Type 2, procédures stockées, reporting), en validant le modèle directement avec le VP Finances.',
        'J\u2019ai reconstruit la couche de données d\u2019un secteur d\u2019affichage extérieur en chargement modulaire, améliorant le temps de rafraîchissement des rapports de 45 %.',
        'J\u2019ai reconstruit 18 rapports hérités et livré 13 nouveaux rapports à travers la migration et les secteurs d\u2019affaires.',
        'Sur un mandat en aviation de 6 mois, j\u2019ai livré 24 rapports Power BI et encadré 5 collègues; sur un mandat en santé publique de 8 mois, j\u2019ai formé 6 départements à bâtir et maintenir leurs propres rapports.',
    ]
)

add_role(
    'CGI',
    'Data Analyst',
    'Montréal, QC',
    'Juillet 2021 - Août 2022',
    [
        'J\u2019ai mis en place l\u2019exécution en parallèle des paquets SSIS, réduisant la fenêtre de chargement globale de 45 %.',
        'J\u2019ai revu les processus ETL Pentaho existants pour éliminer les goulots d\u2019étranglement et permettre le parallélisme, en amont d\u2019une migration vers Talend.',
        'J\u2019ai bâti le premier référentiel de documentation de la plateforme, couvrant ~150 processus ETL et cartographiant leurs dépendances d\u2019exécution.',
        'J\u2019ai mené des ateliers de découverte avec 8 fonctions d\u2019affaires et conçu des modèles de données logiques et physiques.',
        'J\u2019ai contribué au déploiement des politiques et normes de gestion des données à l\u2019échelle de l\u2019entreprise.',
    ]
)

add_role(
    'AECSA',
    'Directeur de Données Stratégique',
    'Bogotá, Colombie',
    'Avril 2017 - Juillet 2018',
    [
        'J\u2019ai dirigé la conception d\u2019un nouvel écosystème de données centralisé pour remplacer le reporting fragmenté portefeuille par portefeuille par un seul environnement intégré.',
        'J\u2019ai interrogé 9 directeurs de comptes et défini un ensemble standard de 12 KPI adopté sur les 30 portefeuilles.',
        'J\u2019ai conçu l\u2019architecture autour d\u2019exigences strictes de ségrégation des données entre institutions financières.',
        'J\u2019ai bâti et dirigé une équipe de 6 personnes (3 analystes d\u2019affaires, 3 développeurs BI), en livrant un pilote intégrant 2 portefeuilles.',
        'J\u2019ai bâti l\u2019analyse de rentabilité pour un état cible infonuagique et dirigé la négociation avec Microsoft, incluant la livraison de Power BI sur mobile.',
    ]
)

add_role(
    'Banco Falabella Colombia',
    'Coordinateur de Campagnes – Business Intelligence',
    'Bogotá, Colombie',
    'Décembre 2014 - Novembre 2016',
    [
        'J\u2019ai bâti des profils clients consolidés par modélisation RFM sur 1,5 million de clients, faisant passer le ciblage de la détention de carte au comportement.',
        'J\u2019ai priorisé le pipeline de campagnes par impact attendu par secteur; la facturation a augmenté de 8 % d\u2019une année à l\u2019autre.',
        'J\u2019ai réduit l\u2019abandon des canaux numériques de 14 % en appliquant la réglementation colombienne sur le recontact.',
        'J\u2019ai réduit les coûts de communication de 35 % en déplaçant le volume vers le courriel et standardisé le reporting de campagnes en un seul ensemble de KPI.',
        'J\u2019ai développé le scoring de propension de vente croisée et le moteur de sélection de clients en PHP sur SQL Server.',
    ]
)

add_role(
    'INFOMEDIA / SoftTek Solutions',
    'Chef de Projet TI – Business Intelligence',
    'Mexico, Mexique',
    'Août 2013 - Septembre 2014',
    [
        'J\u2019ai conçu l\u2019extraction et la consolidation des données transactionnelles afin de tracer les mouvements de fonds de bout en bout, à un volume de ~210 millions de transactions par mois.',
        'J\u2019ai déplacé les données du mainframe (DB2) vers Teradata et mis en place une extraction en parallèle qui a amélioré les temps d\u2019extraction de 40 %.',
        'J\u2019ai implémenté le profilage de risque client et la logique d\u2019identification des flux d\u2019argent selon les règles de détection LBA/AML.',
        'J\u2019ai assuré la gouvernance de livraison de ~25 rapports réglementaires interdépendants, en suivant leur état avec des intervenants de niveau VP.',
    ]
)

add_role(
    'BBVA Colombia',
    'Spécialiste BI (progression depuis Développeur COBOL)',
    'Bogotá, Colombie',
    'Janvier 2008 - Juillet 2013',
    [
        'J\u2019ai agi comme leader technique pour la Colombie sur le programme corporatif de Campagnes Commerciales du groupe, déployé ensuite au Pérou, en Argentine et au Chili.',
        'J\u2019ai été responsable du moteur de règles d\u2019affaires calculant la probabilité de chaque client d\u2019accepter une offre de crédit et transmettant les profils qualifiés au Risque.',
        'J\u2019ai automatisé le profilage client et le transfert au Risque sur tous les canaux; l\u2019efficacité des campagnes est passée de 45 % à 66 %.',
        'J\u2019ai bâti une transaction de reporting de risque opérationnel à l\u2019échelle de la banque (COBOL/JCL/CICS) déployée dans 350 succursales.',
    ]
)

# Expérience Additionnelle
add_section_heading('Expérience Additionnelle')
add_plain('Consultant Indépendant en Données et Analytique  |  À distance depuis Montréal, QC  |  Février 2019 - Avril 2020')
add_plain('oXya, A Hitachi Group Company — Technicien des Opérations  |  Montréal, QC  |  Janvier 2021 - Juillet 2021')
add_plain('InteRecycle — Spécialiste du Support Technique  |  Montréal, QC  |  Juillet 2020 - Septembre 2020')

# Compétences Techniques
add_section_heading('Compétences Techniques')
skills = [
    'BI et Analytique : Power BI, Power BI Service, DAX, Power Query (M), SSRS, visualisation de données, BI en libre-service',
    'Données et Ingénierie : SQL / T-SQL, Snowflake, Microsoft Fabric, Microsoft SQL Server, entreposage de données, ETL/ELT, modélisation dimensionnelle (Kimball), SCD Type 2, Python',
    'Qualité et Gouvernance des Données : validation et réconciliation des données, lignage des données, gouvernance, certification de migration',
    'CRM et Analytique Financière : Salesforce Marketing Cloud, analytique de campagnes, segmentation client (RFM), analytique bancaire',
    'Collaboration et Livraison : Jira, Confluence, Git / Bitbucket, Agile / Scrum / Kanban',
    'Leadership : gestion d\u2019intervenants exécutifs, storytelling de données, présentations exécutives, leadership interfonctionnel, gestion du changement',
]
for s in skills:
    add_bullet(s)

# Formation
add_section_heading('Formation')
add_plain('Cours de 2e cycle – Big Data et Business Intelligence — Montréal, Québec, Canada — 2021')
add_plain('Diplôme d\u2019Études Professionnelles (DEP) – Soutien Informatique — Montréal, Québec, Canada — 2018 – 2020')
add_plain('Spécialisation de 2e cycle en Marketing Numérique — 2014 – 2015')
add_plain('Baccalauréat en Génie des Systèmes — 2002 – 2007')

# Langues
add_section_heading('Langues')
add_plain('Espagnol — Langue Maternelle  |  Français — Compétence Professionnelle  |  Anglais — Compétence Professionnelle')

# Sauvegarde
output_path = r'c:\Users\MalFel01\OneDrive - Videotron Ltée\Personal\Git\bi-data-engineering-knowledge-base\11_CV\CV_FR\CV_Felipe_Maldonado_FR.docx'
doc.save(output_path)
print(f'DOCX créé : {output_path}')
